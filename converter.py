import re
import html
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Iterable, Union

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


# =========================
# CONFIG
# =========================

OUTPUT_META_LABELS = ["Title", "Description", "URL", "Territories", "Target Keyword"]

INPUT_KEY_MAP = {
    "title": "Title",
    "seo title": "Title",
    "meta title": "Title",
    "description": "Description",
    "meta description": "Description",
    "url": "URL",
    "territories": "Territories",
    "territory": "Territories",
    "target keyword": "Target Keyword",
    "target keywords": "Target Keyword",
    "kw": "Target Keyword",
    "keyword": "Target Keyword",
    "primary keyword": "Target Keyword",
    "h1": "H1",
}

TESTO_KEYS = {"testo", "content", "article", "body", "testo articolo"}


# =========================
# HTML entities (SOLO OUTPUT)
# =========================

def html_entities(s: str) -> str:
    if not s:
        return ""

    anchors = []

    def stash_anchor(m):
        anchors.append(m.group(0))
        return f"__ANCHOR_{len(anchors)-1}__"

    s = re.sub(r"<a\s+href=\"[^\"]+\">.*?</a>", stash_anchor, s, flags=re.I | re.S)

    s = html.escape(s, quote=False)

    replacements = {
        "’": "&rsquo;",
        "‘": "&lsquo;",
        "“": "&ldquo;",
        "”": "&rdquo;",
        "–": "&ndash;",
        "—": "&mdash;",
        "…": "&hellip;",
        "à": "&agrave;",
        "è": "&egrave;",
        "é": "&eacute;",
        "ì": "&igrave;",
        "ò": "&ograve;",
        "ù": "&ugrave;",
        "À": "&Agrave;",
        "È": "&Egrave;",
        "É": "&Eacute;",
        "Ì": "&Igrave;",
        "Ò": "&Ograve;",
        "Ù": "&Ugrave;",
    }

    for k, v in replacements.items():
        s = s.replace(k, v)

    for i, a in enumerate(anchors):
        s = s.replace(f"__ANCHOR_{i}__", a)

    return s


# =========================
# DOCX helpers
# =========================

def shade_cell(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)

def set_cell_text(cell, text: str, bold=False, color=None, size_pt=10):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text or "")
    r.bold = bold
    r.font.size = Pt(size_pt)
    if color:
        r.font.color.rgb = color


# =========================
# Iterazione DOCX
# =========================

def iter_block_items(parent) -> Iterable[Union[Paragraph, Table]]:
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl

    parent_elm = parent.element.body if hasattr(parent, "element") else parent._tc
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# =========================
# Hyperlink-aware paragraph
# =========================

def _iter_text_runs(node) -> str:
    texts = []
    for t in node.findall(".//w:t", namespaces=node.nsmap):
        if t.text:
            texts.append(t.text)
    return "".join(texts)
def run_to_html(run_elm) -> str:
    """
    Converte un singolo <w:r> in HTML,
    preservando il grassetto come <strong>.
    """
    texts = []
    is_bold = False

    # verifica grassetto
    rpr = run_elm.find(".//w:rPr", namespaces=run_elm.nsmap)
    if rpr is not None and rpr.find(".//w:b", namespaces=run_elm.nsmap) is not None:
        is_bold = True

    # testo del run
    for t in run_elm.findall(".//w:t", namespaces=run_elm.nsmap):
        if t.text:
            texts.append(t.text)

    text = "".join(texts)
    if not text.strip():
        return ""

    if is_bold:
        return f"<strong>{text}</strong>"

    return text

def paragraph_to_text_with_links(paragraph: Paragraph) -> str:
    out = []
    part = paragraph.part

    for child in paragraph._p.iterchildren():
        tag = child.tag

        # hyperlink
        if tag.endswith("}hyperlink"):
            r_id = child.get(qn("r:id"))
            text = _iter_text_runs(child)

            if r_id and r_id in part.rels and text.strip():
                href = part.rels[r_id].target_ref
                out.append(f'<a href="{href}">{text}</a>')

        # run normale (NON dentro hyperlink)
        elif tag.endswith("}r"):
            if child.getparent() is not paragraph._p:
                continue

            text = _iter_text_runs(child)
            if text.strip():
                out.append(text)

    return "".join(out).strip()


# =========================
# Estrazione RAW
# =========================

def extract_lines_raw(doc: Document) -> List[str]:
    lines = []

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            t = paragraph_to_text_with_links(block)
            if t:
                lines.append(t)

        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t = paragraph_to_text_with_links(p)
                        if t:
                            lines.append(t)

    return lines


# =========================
# Parsing input DOCX
# =========================

def parse_input_docx(path: Path) -> Dict[str, Any]:
    doc = Document(str(path))
    lines = extract_lines_raw(doc)

    meta = {k: "" for k in OUTPUT_META_LABELS}
    body: List[Dict[str, str]] = []
    h1 = ""

    in_testo = False
    testo_re = re.compile(r"^({})\s*:\s*$".format("|".join(TESTO_KEYS)), re.I)

    for line in lines:
        if testo_re.match(line):
            in_testo = True
            continue

        if not in_testo:
            m = re.match(r"^([^:]{1,80})\s*:\s*(.*)$", line)
            if m:
                key = m.group(1).strip().lower()
                val = m.group(2).strip()
                if key in INPUT_KEY_MAP:
                    out = INPUT_KEY_MAP[key]
                    if out == "H1":
                        h1 = val
                    elif out in meta:
                        meta[out] = val
            continue

        m = re.match(r"^(.*)\s*\((h2|h3)\)\s*$", line, re.I)
        if m:
            body.append({
                "block": "✏️ S3",
                "html": f"<h2><strong>{html_entities(m.group(1))}</strong></h2>"
            })
        else:
            body.append({
                "block": "Intro" if not body else "✏️ S3",
                "html": f'<p class="h-text-size-14 h-font-primary">{html_entities(line)}</p>'
            })

    if not h1:
        h1 = meta.get("Title") or "Untitled"

    return {
        "meta": meta,
        "h1": h1,
        "body": body
    }


# =========================
# Structure of content
# =========================

def build_structure(body: List[Dict[str, str]]) -> List[str]:
    s = ["H1", "Intro"]
    s3 = sum(1 for b in body if b["block"] == "✏️ S3" and b["html"].startswith("<h2"))
    s.extend(["✏️ S3"] * s3)
    return s


# =========================
# Output DOCX
# =========================

def write_output_docx(parsed: Dict[str, Any], out: Path):
    doc = Document()
    meta = parsed["meta"]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta.get("Title") or parsed["h1"])
    r.bold = True
    r.font.size = Pt(20)

    doc.add_paragraph("")
    doc.add_paragraph("")

    table = doc.add_table(rows=len(OUTPUT_META_LABELS), cols=2)
    table.style = "Table Grid"

    for i, k in enumerate(OUTPUT_META_LABELS):
        shade_cell(table.cell(i, 0), "000000")
        set_cell_text(table.cell(i, 0), k, bold=True, color=RGBColor(255, 255, 255))
        set_cell_text(table.cell(i, 1), meta.get(k, ""))

    doc.add_paragraph("")
    doc.add_paragraph("")

    p = doc.add_paragraph("Structure of content:")
    p.runs[0].bold = True
    for l in build_structure(parsed["body"]):
        doc.add_paragraph(l)

    doc.add_paragraph("")
    doc.add_paragraph("")

    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    set_cell_text(t2.cell(0, 0), "Block", bold=True)
    set_cell_text(t2.cell(0, 1), "⭐ HTML Output ⭐", bold=True)

    for item in parsed["body"]:
        row = t2.add_row().cells
        row[0].text = item["block"]
        row[1].add_paragraph(item["html"])

    doc.save(str(out))


# =========================
# Streamlit helper
# =========================

def convert_uploaded_file(uploaded_file):
    with tempfile.TemporaryDirectory() as d:
        d = Path(d)
        inp = d / uploaded_file.name
        out = d / f"output_{uploaded_file.name}"
        inp.write_bytes(uploaded_file.read())
        parsed = parse_input_docx(inp)
        write_output_docx(parsed, out)
        final = Path(tempfile.gettempdir()) / out.name
        final.write_bytes(out.read_bytes())
        return final

