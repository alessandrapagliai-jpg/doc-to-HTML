import streamlit as st
import re
import html
import tempfile
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# =========================
# CONFIGURAZIONE E MAPPE
# =========================
OUTPUT_META_LABELS = ["Title", "Description", "URL", "Territories", "Target Keyword"]
INPUT_KEY_MAP = {
    "title": "Title", "seo title": "Title", "meta title": "Title",
    "description": "Description", "meta description": "Description",
    "url": "URL", "territories": "Territories",
    "target keyword": "Target Keyword", "kw": "Target Keyword", "keyword": "Target Keyword",
    "h1": "H1",
}

# =========================
# FUNZIONI DI CONVERSIONE
# =========================

def get_html_text(paragraph) -> str:
    """Estrae il testo preservando il grassetto (strong)."""
    full_html = ""
    for run in paragraph.runs:
        text = html.escape(run.text)
        if run.bold:
            full_html += f"<strong>{text}</strong>"
        else:
            full_html += text
    
    replacements = {"’": "&rsquo;", "à": "&agrave;", "è": "&egrave;", "é": "&eacute;", "ì": "&igrave;", "ò": "&ograve;", "ù": "&ugrave;"}
    for k, v in replacements.items():
        full_html = full_html.replace(k, v)
    return full_html

def parse_input_docx(path: Path):
    doc = Document(str(path))
    meta = {k: "" for k in OUTPUT_META_LABELS}
    h1 = ""
    body_elements = []
    
    # 1. Metadati dalle Tabelle [cite: 2]
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip().lower().replace(":", "")
                val = row.cells[1].text.strip()
                if key in INPUT_KEY_MAP:
                    mapped_key = INPUT_KEY_MAP[key]
                    if mapped_key == "H1":
                        h1 = val
                    else:
                        meta[mapped_key] = val

    # 2. Contenuto e Tag (h2/h3) [cite: 2]
    for para in doc.paragraphs:
        text_raw = para.text.strip()
        if not text_raw or text_raw.lower().startswith("testo:"):
            continue
        
        header_match = re.search(r"\((h2|h3)\)$", text_raw, re.I)
        if header_match:
            tag = header_match.group(1).lower()
            clean_text = re.sub(r"\s*\(h[23]\)$", "", text_raw, flags=re.I)
            body_elements.append({"block": "✏️ S3", "html": f"<{tag}><strong>{clean_text}</strong></{tag}>"})
        else:
            html_content = get_html_text(para)
            block_name = "Intro" if not any(b["block"] == "✏️ S3" for b in body_elements) else "✏️ S3"
            body_elements.append({"block": block_name, "html": f'<p class="h-text-size-14 h-font-primary">{html_content}</p>'})

    if not h1: h1 = meta.get("Title", "Untitled")
    return {"meta": meta, "h1": h1, "body": body_elements}

def write_output_docx(parsed, out_path):
    doc = Document()
    
    # Header e Tabella Meta [cite: 2]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(parsed["h1"])
    r.bold, r.font.size = True, Pt(18)

    table = doc.add_table(rows=len(OUTPUT_META_LABELS), cols=2)
    table.style = "Table Grid"
    for i, k in enumerate(OUTPUT_META_LABELS):
        cell = table.cell(i, 0)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "000000")
        tc_pr.append(shd)
        rk = cell.paragraphs[0].add_run(k)
        rk.bold, rk.font.color.rgb = True, RGBColor(255, 255, 255)
        table.cell(i, 1).text = parsed["meta"].get(k, "")

    # Tabella HTML [cite: 2]
    doc.add_paragraph("\n⭐ HTML Output ⭐")
    t_html = doc.add_table(rows=1, cols=2)
    t_html.style = "Table Grid"
    t_html.cell(0,0).text, t_html.cell(0,1).text = "Block", "HTML Code"
    
    for item in parsed["body"]:
        row = t_html.add_row().cells
        row[0].text, row[1].text = item["block"], item["html"]

    doc.save(str(out_path))

def convert_uploaded_file(uploaded_file):
    with tempfile.TemporaryDirectory() as d:
        d_path = Path(d)
        inp = d_path / uploaded_file.name
        out = d_path / f"output_{uploaded_file.name}"
        inp.write_bytes(uploaded_file.read())
        parsed = parse_input_docx(inp)
        write_output_docx(parsed, out)
        # Salvataggio persistente per il download
        final_path = Path(tempfile.gettempdir()) / out.name
        final_path.write_bytes(out.read_bytes())
        return final_path

# =========================
# INTERFACCIA STREAMLIT
# =========================
st.set_page_config(page_title="DOCX → HTML SEO Converter", layout="centered")
st.title("📄 DOCX → HTML SEO Converter")
st.write("Carica i file DOCX per convertirli in HTML SEO-ready.")

uploaded_files = st.file_uploader("Upload file DOCX", type=["docx"], accept_multiple_files=True)

if uploaded_files:
    if st.button("🚀 Converti"):
        with st.spinner("Conversione in corso..."):
            for uf in uploaded_files:
                try:
                    out_path = convert_uploaded_file(uf)
                    with open(out_path, "rb") as f:
                        st.download_button(
                            label=f"⬇️ Scarica {uf.name}",
                            data=f,
                            file_name=f"CONVERTITO_{uf.name}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"Errore su {uf.name}: {e}")
        st.success("Conversione completata!")
